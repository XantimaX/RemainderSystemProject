import pymupdf
from docx import Document
from ComponentExtraction import extract_details
def handout_reader_pdf(pdf_file) :
    details = []

    doc = pymupdf.open(pdf_file)

    course_name = None
    words = doc[0].get_text("words")

    word_contents = ""
    for word in words :
        if (word[4].lower().replace("-", " ") == "instructor in charge") :
            break
        word_contents += word[4]+" "

    course_name = word_contents.split(":")[-1].strip()

    tables = []

    for page in doc :
        if page.search_for("Duration"):
            tables.extend(page.find_tables(strategy = "lines_strict").tables)
            break

    correct_table_location = -1

    for i in range(len(tables)) :
        names = tables[i].header.names
        for name in names :
            if name  and "duration" in name.lower() :
                correct_table_location = i
                break
        else :
            continue
        break

    components = tables[correct_table_location].to_pandas()
    cols = components.columns
    position_of_date = -1
    position_of_evaluation = -1

    for i in range(len(cols)):
        if (position_of_evaluation != -1 and position_of_date != -1) :
            break
        elif "Date" in cols[i] :
            position_of_date = i
        elif "Evaluation" in cols[i] :
            position_of_evaluation = i

    print(course_name)
    components = components.to_numpy()
    for component in components :
        component_name = component[position_of_evaluation]
        component_date = component[position_of_date]
        component_details = extract_details(component_name=component_name, component_date=component_date)

        if (component_details) :
            details.append([course_name, component_details[0], component_details[1]])

    return details

def handout_reader_word (word_file) :
    details = []
    doc = Document(word_file)

    course_name = None

    for p_no,paragrah in enumerate(doc.paragraphs):

        if "Course Title" in paragrah.text :
            course_name = paragrah.text.split(":")[-1].strip()
            break

    correct_table_no = -1
    for table_no,table in enumerate(doc.tables):
        for cell in table.rows[0].cells :
            if "duration" in cell.text.lower() :
                correct_table_no = table_no
                break
        else :
            continue
        break

    correct_table = doc.tables[correct_table_no]
    date_location = -1
    component_location = -1

    for i,cell in enumerate(correct_table.rows[0].cells) :
        if date_location != -1 and component_location != -1 :
            break
        elif cell and "date" in cell.text.lower() :
            date_location = i
        elif cell and "components" in cell.text.lower() :
            component_location = i

    print(course_name)
    for row in correct_table.rows[1:] :
        component_name = row.cells[component_location].text
        component_date = row.cells[date_location].text
        component_details = extract_details(component_name=component_name, component_date=component_date)
        if (component_details) :
            details.append([course_name, component_details[0], component_details[1]])

    return details
